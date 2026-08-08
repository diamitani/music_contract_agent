#!/usr/bin/env python3
"""
Portable contract renderer — no dependency on Claude's bundled docx skill,
OpenCode tools, or any other harness-specific capability. Any agent (or a
human) can call this directly with `python3` as long as `python-docx` is
installed (`pip install python-docx`).

This is a FALLBACK, not the preferred path. If the environment running this
skill provides a richer, harness-native document-generation capability
(e.g. Claude's docx skill, an OpenCode file-writing tool, a Word COM
automation library), prefer that — it will produce better typography,
tables, and page layout. Use this script only when no such capability is
available, or when you need a fully scriptable, dependency-light path
(e.g. running this library in a plain CI job or a headless server).

Input: a JSON file describing the contract (see --help / SPEC below).
Output: a .docx file.

Usage:
    python3 scripts/render_docx.py --input contract.json --output out.docx

JSON input shape:
{
  "title": "PRODUCER AGREEMENT",
  "effective_date": "August 8, 2026",
  "parties": ["Nova Ray (\"Artist\")", "Kid Static (\"Producer\")"],
  "sections": [
    {"heading": "1. Services", "body": "Producer will produce and mix..."},
    {"heading": "2. Compensation", "body": "Artist will pay Producer..."}
  ],
  "signature_blocks": ["Artist", "Producer"],
  "notice": "Important: This document is an artist-support drafting tool..."
}
"""
import argparse
import json
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(
        "python-docx is not installed. Install it with:\n"
        "    pip install python-docx\n"
        "(or --break-system-packages if your environment requires it)",
        file=sys.stderr,
    )
    sys.exit(1)


def build_docx(spec: dict, output_path: str):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    title = doc.add_heading(spec.get("title", "AGREEMENT"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if spec.get("effective_date"):
        p = doc.add_paragraph(f"Effective Date: {spec['effective_date']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if spec.get("parties"):
        p = doc.add_paragraph()
        p.add_run("Between: ").bold = True
        p.add_run(" and ".join(spec["parties"]))

    doc.add_paragraph()

    for sec in spec.get("sections", []):
        heading = sec.get("heading", "")
        body = sec.get("body", "")
        if heading:
            h = doc.add_heading(heading, level=2)
            for run in h.runs:
                run.font.size = Pt(12)
        for para in body.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    blocks = spec.get("signature_blocks", [])
    if blocks:
        doc.add_page_break()
        doc.add_heading("Signatures", level=2)
        table = doc.add_table(rows=len(blocks) * 3, cols=1)
        table.autofit = True
        row_i = 0
        for role in blocks:
            table.cell(row_i, 0).text = f"{role}: ___________________________"
            row_i += 1
            table.cell(row_i, 0).text = "Print Name: ___________________________"
            row_i += 1
            table.cell(row_i, 0).text = "Date: ___________________________"
            row_i += 1
            row_i += 0

    if spec.get("notice"):
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run(spec["notice"])
        run.italic = True
        run.font.size = Pt(9)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Render a contract JSON spec to .docx")
    parser.add_argument("--input", required=True, help="Path to input JSON spec")
    parser.add_argument("--output", required=True, help="Path to write the .docx file")
    args = parser.parse_args()

    with open(args.input, encoding="utf-8") as f:
        spec = json.load(f)

    build_docx(spec, args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
